# -*- coding: utf-8 -*-
"""
Plantilla de CV en .docx (ATS-friendly + moderno con acento de color)
Uso: python cv_template/generate_cv.py
Requiere: python-docx (pip install python-docx)
Pillow es OPCIONAL (solo si quieres recortar/estandarizar la foto antes de insertarla)

Salida: output/CV_<Nombre>_<Apellido>_<Año>.docx
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


# -----------------------------
# Configuración visual
# -----------------------------
PRIMARY_HEX = "2F80ED"        # Acento profesional (azul petróleo)
FONT_FAMILY = "Calibri"       # ATS-friendly
MARGIN_CM = 2.0               # Márgenes simétricos (ATS + legible)
PHOTO_WIDTH_CM = 3.5          # Ancho visual sugerido para foto en el documento


# -----------------------------
# Utilidades de estilo
# -----------------------------
def set_margins(section, left=MARGIN_CM, right=MARGIN_CM, top=MARGIN_CM, bottom=MARGIN_CM):
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)


def set_run_style(run, size=11, bold=False, color_rgb=None, font_family=FONT_FAMILY):
    run.font.name = font_family
    # Asegurar font en EastAsia para MS Word
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
    run.font.size = Pt(size)
    run.bold = bold
    if color_rgb:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*color_rgb)


def _hex_to_rgb_tuple(hexstr):
    hexstr = hexstr.replace("#", "")
    return tuple(int(hexstr[i:i+2], 16) for i in (0, 2, 4))


def add_colored_divider(paragraph, color_hex=PRIMARY_HEX, size=6, space=1):
    """
    Inserta una línea divisoria (borde inferior) al párrafo, sin usar tablas (mejor para ATS).
    size en octavos de punto (Word). 6 ≈ 0.75pt
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color_hex.replace("#", ""))
    pBdr.append(bottom)

    pPr.append(pBdr)


def add_section_title(doc, text, color_hex=PRIMARY_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_style(run, size=12, bold=True)
    # Línea de acento debajo del título
    add_colored_divider(p, color_hex=color_hex, size=10, space=5)


def add_bullets(doc, items, left_indent_cm=0.5, space_after_pt=2):
    for it in items or []:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(left_indent_cm)
        p.paragraph_format.space_after = Pt(space_after_pt)
        run = p.add_run(it)
        set_run_style(run, size=10.5)


def _add_header_text_block(container, nombre, cargo, contacto, ubicacion=None):
    # contenedor es un Document o una celda de tabla (tiene add_paragraph)
    p_name = container.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_style(p_name.add_run(nombre), size=20, bold=True)

    p_role = container.add_paragraph()
    set_run_style(p_role.add_run(cargo), size=12, bold=False, color_rgb=_hex_to_rgb_tuple(PRIMARY_HEX))

    # Filtrar LinkedIn de los datos de contacto
    def _is_linkedin(s: str) -> bool:
        return 'linkedin' in s.lower()

    contacts_clean = [c for c in (contacto or []) if c and not _is_linkedin(c)]
    contact_line = " | ".join(contacts_clean)  # email | móvil | web (sin LinkedIn)
    if ubicacion:
        contact_line = f"{contact_line} | {ubicacion}" if contact_line else ubicacion
    p_contact = container.add_paragraph(contact_line)
    for run in p_contact.runs:
        set_run_style(run, size=10.5)


def add_header(doc, nombre, cargo, contacto, ubicacion=None, incluir_foto=False, ruta_foto=None, photo_position='right_paragraph'):
    """
    photo_position:
      - 'right_paragraph' (por defecto): foto en párrafo independiente alineado a la derecha.
      - 'right_table': texto a la izquierda, foto a la derecha (encabezado en tabla 2 columnas).
      - 'left_table': foto a la izquierda, texto a la derecha (encabezado en tabla 2 columnas).
    """
    use_table = photo_position in ('right_table', 'left_table')

    if incluir_foto and ruta_foto and os.path.exists(ruta_foto) and use_table:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]

        if photo_position == 'right_table':
            _add_header_text_block(left_cell, nombre, cargo, contacto, ubicacion)
            p_photo = right_cell.paragraphs[0]
            p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                p_photo.add_run().add_picture(ruta_foto, width=Cm(PHOTO_WIDTH_CM))
            except Exception as ex:
                err = right_cell.add_paragraph(f"(No se pudo insertar la foto: {ex})")
                for r in err.runs: set_run_style(r, size=9)
        else:  # left_table
            p_photo = left_cell.paragraphs[0]
            p_photo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                p_photo.add_run().add_picture(ruta_foto, width=Cm(PHOTO_WIDTH_CM))
            except Exception as ex:
                err = left_cell.add_paragraph(f"(No se pudo insertar la foto: {ex})")
                for r in err.runs: set_run_style(r, size=9)
            _add_header_text_block(right_cell, nombre, cargo, contacto, ubicacion)
    else:
        # Encabezado sin tabla (ATS-friendly): texto y, opcionalmente, foto en párrafo derecho
        _add_header_text_block(doc, nombre, cargo, contacto, ubicacion)
        if incluir_foto and ruta_foto and os.path.exists(ruta_foto):
            p_photo = doc.add_paragraph()
            p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                p_photo.add_run().add_picture(ruta_foto, width=Cm(PHOTO_WIDTH_CM))
            except Exception as ex:
                p_err = doc.add_paragraph(f"(No se pudo insertar la foto: {ex})")
                for run in p_err.runs:
                    set_run_style(run, size=9)


def add_resumen(doc, texto):
    add_section_title(doc, "Resumen profesional")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    set_run_style(run, size=11)


def add_experiencia(doc, experiencias):
    add_section_title(doc, "Experiencia profesional")
    for exp in experiencias or []:
        # Encabezado del puesto
        header = f"{exp['puesto']} – {exp['empresa']} | {exp['periodo']}"
        p = doc.add_paragraph()
        set_run_style(p.add_run(header), size=11.5, bold=True)

        # Subtítulo (opcional)
        subt = []
        if exp.get('ubicacion'): subt.append(exp['ubicacion'])
        if exp.get('sector'): subt.append(exp['sector'])
        if subt:
            p2 = doc.add_paragraph(" / ".join(subt))
            for r in p2.runs: set_run_style(r, size=10)

        # Bloques
        if exp.get('logros'):
            p3 = doc.add_paragraph()
            set_run_style(p3.add_run("Logros:"), size=10.5, bold=True)
            add_bullets(doc, exp['logros'])
        if exp.get('actividades'):
            p4 = doc.add_paragraph()
            set_run_style(p4.add_run("Actividades:"), size=10.5, bold=True)
            add_bullets(doc, exp['actividades'])
        if exp.get('proyectos'):
            p5 = doc.add_paragraph()
            set_run_style(p5.add_run("Proyectos:"), size=10.5, bold=True)
            add_bullets(doc, exp['proyectos'])


def add_educacion(doc, items):
    add_section_title(doc, "Educación")
    for ed in items or []:
        p = doc.add_paragraph()
        set_run_style(p.add_run(f"{ed['grado']} – {ed['institucion']}"), size=11.5, bold=True)
        if ed.get('detalle'):
            p2 = doc.add_paragraph(ed['detalle'])
            for r in p2.runs: set_run_style(r, size=10.5)


def add_certificaciones(doc, certs):
    if not certs: return
    add_section_title(doc, "Certificaciones")
    add_bullets(doc, certs, left_indent_cm=0.4)


def add_habilidades(doc, skills):
    add_section_title(doc, "Habilidades")
    add_bullets(doc, skills, left_indent_cm=0.4)


def add_idiomas(doc, idiomas):
    add_section_title(doc, "Idiomas")
    pairs = [f"{k}: {v}" for k, v in (idiomas or {}).items()]
    add_bullets(doc, pairs, left_indent_cm=0.4)


# -----------------------------
# Datos de ejemplo (puedes cargar desde JSON)
# -----------------------------
def demo_data():
    return {
        "nombre": "Natalia Moreno",
        "cargo": "Consultora SAP CPI",
        "contacto": ["natalia.moreno9112@gmail.com", "+51 999 888 777", "linkedin.com/in/nmoreno"],
        "ubicacion": "Lima, Perú",
        "incluir_foto": True,
        "ruta_foto": os.path.join(os.path.dirname(__file__), "assets", "photos", "nataliamorenocv.png"),
        "photo_position": "right_table",
        "resumen": (
            "Consultora SAP CPI con más de 5 años en integraciones SAP ↔ sistemas externos. "
            "Enfoque en rendimiento, seguridad y continuidad operativa. Especialista en iFlows, "
            "automatización y monitoreo proactivo con impacto medible en SLAs."
        ),
        "experiencias": [
            {
                "puesto": "Consultora SAP CPI",
                "empresa": "EXOLMAR S.A.A.",
                "periodo": "2022 – Actualidad",
                "ubicacion": "Remoto / Lima",
                "sector": "Industria",
                "logros": [
                    "Reducción del tiempo de integración en ~40% mediante optimización de iFlows.",
                    "Monitoreo proactivo que disminuyó incidentes críticos en ~30%."
                ],
                "actividades": [
                    "Orquestación y versionado de artefactos CPI.",
                    "Soporte L2/L3 y resolución de incidencias.",
                    "Capacitación y handover a operación."
                ],
                "proyectos": [
                    "Integración SAP ↔ CRM (Salesforce).",
                    "Migración de interfaces desde PI/PO hacia CPI."
                ]
            },
            {
                "puesto": "Consultora SAP CPI",
                "empresa": "NEXIS CONSULTING S.A.C",
                "periodo": "2019 – 2022",
                "logros": [
                    "Mejora de performance y reducción de errores en ~50%.",
                    "Diseño de arquitectura escalable con estándares de seguridad."
                ],
                "actividades": [
                    "Desarrollo de interfaces SAP con sistemas externos.",
                    "Implementación de políticas de seguridad y cifrado.",
                    "Pruebas técnicas/funcionales y documentación."
                ],
                "proyectos": [
                    "Plataforma e-commerce retail con CPI.",
                    "Automatización de facturación electrónica con SUNAT."
                ]
            }
        ],
        "educacion": [
            {"grado": "Ingeniería de Sistemas", "institucion": "Pontificia Universidad Católica del Perú"},
            {"grado": "Certificación SAP CPI", "institucion": "SAP SE"}
        ],
        "certificaciones": ["SAP Certified Associate – SAP Integration Suite (CPI)"],
        "habilidades": [
            "SAP CPI / PI-PO",
            "Integraciones B2B/B2C",
            "Facturación electrónica",
            "Seguridad en middleware",
            "Monitoreo y troubleshooting"
        ],
        "idiomas": {"Español": "Nativo", "Inglés": "Intermedio-Avanzado"}
    }


# -----------------------------
# Generación del documento
# -----------------------------
def construir_cv(data, carpeta_salida="output"):
    os.makedirs(carpeta_salida, exist_ok=True)

    doc = Document()
    set_margins(doc.sections[0])

    add_header(
        doc,
        nombre=data["nombre"],
        cargo=data["cargo"],
        contacto=data["contacto"],
        ubicacion=data.get("ubicacion"),
        incluir_foto=data.get("incluir_foto", False),
        ruta_foto=data.get("ruta_foto"),
        photo_position=data.get("photo_position", 'right_paragraph')
    )

    add_resumen(doc, data["resumen"])
    add_experiencia(doc, data["experiencias"])
    add_educacion(doc, data["educacion"])
    add_certificaciones(doc, data.get("certificaciones", []))
    add_habilidades(doc, data["habilidades"])
    add_idiomas(doc, data["idiomas"])

    safe_name = data["nombre"].strip().replace(" ", "_")
    out_name = f"CV_{safe_name}_{datetime.now().year}.docx"
    out_path = os.path.join(carpeta_salida, out_name)
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    data = demo_data()
    construir_cv(data)
